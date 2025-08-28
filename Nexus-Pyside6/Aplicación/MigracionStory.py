import sys
import os
import time
import threading
import google.generativeai as genai
import docx
from pypdf import PdfReader
from PySide6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QTextEdit, QPushButton, QFileDialog, QComboBox, QProgressBar,
    QMessageBox, QFrame, QPlainTextEdit, QStackedWidget
)
from PySide6.QtCore import Qt, Signal, QObject, QThread, Qt
from PySide6.QtGui import QIcon
import re


# -----------------------------
# Funciones auxiliares
# -----------------------------
def extract_text_from_file(file_path):
    """Extrae texto de archivos .docx o .pdf."""
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
    else:
        raise ValueError("Formato de archivo no soportado. Usa .docx o .pdf.")


def split_document_into_chunks(text, max_chunk_size=3000):
    """Divide el documento en chunks manejables."""
    # Primero intentar dividir por secciones/capítulos
    sections = re.split(r'\n\s*(?:[0-9]+\.|\b(?:CAPÍTULO|SECCIÓN|MÓDULO|FUNCIONALIDAD)\b)', text, flags=re.IGNORECASE)

    chunks = []
    current_chunk = ""

    for section in sections:
        if len(current_chunk) + len(section) < max_chunk_size:
            current_chunk += section
        else:
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            current_chunk = section

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Si no hay divisiones claras, dividir por párrafos
    if len(chunks) == 1 and len(text) > max_chunk_size:
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""

        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

    return chunks


def create_analysis_prompt(document_text, role):
    """Crea un prompt inicial para análisis de funcionalidades."""
    return f"""
Eres un analista de negocios Senior. Tu tarea es IDENTIFICAR Y LISTAR todas las funcionalidades del siguiente documento.

DOCUMENTO A ANALIZAR:
{document_text}

INSTRUCCIONES:
1. Lee COMPLETAMENTE el documento
2. Identifica TODAS las funcionalidades mencionadas
3. Crea una LISTA NUMERADA de funcionalidades EXCLUSIVAMENTE para el rol: {role}.
4. Ignora cualquier funcionalidad que corresponda a otros roles diferentes a {role}.

FORMATO DE RESPUESTA:
Lista de Funcionalidades Identificadas:
1. [Nombre funcionalidad] - [Descripción breve]
2. [Nombre funcionalidad] - [Descripción breve]
...

Al final indica: "TOTAL FUNCIONALIDADES IDENTIFICADAS: [número]"

NO generes historias de usuario todavía, solo la lista de funcionalidades.
"""


def create_story_generation_prompt(functionalities_list, document_text, role, start_index, batch_size=5):
    """Crea prompt para generar historias de usuario por lotes."""
    end_index = min(start_index + batch_size, len(functionalities_list))
    selected_functionalities = functionalities_list[start_index:end_index]

    func_text = "\n".join([f"{i + start_index + 1}. {func}" for i, func in enumerate(selected_functionalities)])

    return f"""
Eres un analista de negocios Senior. Genera historias de usuario DETALLADAS para las siguientes funcionalidades específicas.

FUNCIONALIDADES A DESARROLLAR (Lote {start_index + 1} a {end_index}):
{func_text}

DOCUMENTO DE REFERENCIA (para contexto adicional):
{document_text[:2000]}...

FORMATO OBLIGATORIO para CADA funcionalidad:

```
═══════════════════════════════════════════════════════════════
HISTORIA #{start_index + 1}: [Título de la funcionalidad]
═══════════════════════════════════════════════════════════════

COMO: {role}
QUIERO: [funcionalidad específica y detallada]
PARA: [beneficio de negocio claro y medible]

CRITERIOS DE ACEPTACIÓN:

🔹 Escenario Principal:
   DADO que [contexto específico]
   CUANDO [acción concreta del usuario]
   ENTONCES [resultado esperado detallado]

🔹 Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acción diferente]
   ENTONCES [resultado alternativo]

🔹 Validaciones:
   DADO que [condición de error]
   CUANDO [acción que genera error]
   ENTONCES [manejo de error esperado]

REGLAS DE NEGOCIO:
• [Regla específica 1]
• [Regla específica 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

═══════════════════════════════════════════════════════════════
```

IMPORTANTE: TODAS las historias deben generarse ÚNICAMENTE desde la perspectiva del rol **{role}**.
No inventes ni incluyas otros roles diferentes a {role}.
Numera consecutivamente desde {start_index + 1}.
"""


def create_prompt(document_text, role, story_type):
    """Mantiene compatibilidad con el método original pero mejorado."""

    if story_type == 'HU Usuario Funcional':
        # Para documentos grandes, usar estrategia de chunks
        if len(document_text) > 5000:
            return "CHUNK_PROCESSING_NEEDED"

        # Para documentos medianos/pequeños, prompt optimizado
        prompt = f"""
Eres un analista de negocios Senior especializado en QA y análisis exhaustivo de requerimientos.

DOCUMENTO A ANALIZAR:
{document_text}

INSTRUCCIONES CRÍTICAS:

1. ANÁLISIS EXHAUSTIVO:
   - Identifica TODAS las funcionalidades del documento
   - Incluye ÚNICAMENTE las que correspondan al rol que se proporciona en la UI {role}

2. GENERACIÓN DE HISTORIAS PARA: **{role}**

FORMATO OBLIGATORIO:

```
═══════════════════════════════════════════════════════════════
HISTORIA #{{número}}: [Título]
═══════════════════════════════════════════════════════════════

COMO: {role}
QUIERO: [funcionalidad específica]
PARA: [beneficio de negocio]

CRITERIOS DE ACEPTACIÓN:

🔹 Escenario Principal:
   DADO que [contexto]
   CUANDO [acción]
   ENTONCES [resultado]

🔹 Escenario Alternativo:
   DADO que [contexto alternativo]
   CUANDO [acción diferente]
   ENTONCES [resultado alternativo]

🔹 Validaciones:
   DADO que [error]
   CUANDO [acción error]
   ENTONCES [manejo error]

REGLAS DE NEGOCIO:
• [Regla 1]
• [Regla 2]

PRIORIDAD: [Alta/Media/Baja]
COMPLEJIDAD: [Simple/Moderada/Compleja]

═══════════════════════════════════════════════════════════════
```

EXPECTATIVA: Genera entre 10-50 historias según el contenido del documento.

IMPORTANTE: Si el documento es extenso y sientes que podrías cortarte, termina la historia actual y agrega al final:
"CONTINÚA EN EL SIGUIENTE LOTE - FUNCIONALIDADES PENDIENTES: [lista las que faltan]"
"""

    elif story_type == 'HU no Funcionales':
        prompt = f"""
Eres un analista de negocios Senior especializado en requisitos no funcionales.

DOCUMENTO A ANALIZAR:
{document_text}

Identifica TODOS los requisitos no funcionales (rendimiento, seguridad, usabilidad, etc.) y genera historias para el rol: {role}

FORMATO:

```
═══════════════════════════════════════════════════════════════
HISTORIA NO FUNCIONAL #{{número}}: [Título]
═══════════════════════════════════════════════════════════════

COMO: {role}
NECESITO: [requisito no funcional]
PARA: [garantizar calidad]

CRITERIOS DE ACEPTACIÓN:
• [Criterio medible 1]
• [Criterio medible 2]

MÉTRICAS:
• [Métrica objetivo]

CATEGORÍA: [Rendimiento/Seguridad/Usabilidad/etc.]
PRIORIDAD: [Alta/Media/Baja]

═══════════════════════════════════════════════════════════════
```
"""

    elif story_type == 'Requisitos técnicos':
        prompt = f"""
Eres un arquitecto de software Senior.

DOCUMENTO A ANALIZAR:
{document_text}

Identifica TODOS los requisitos técnicos y genera especificaciones detalladas.

FORMATO:

```
═══════════════════════════════════════════════════════════════
REQUISITO TÉCNICO #{{número}}: [Título]
═══════════════════════════════════════════════════════════════

DESCRIPCIÓN:
[Descripción detallada]

CONSIDERACIONES TÉCNICAS:
• [Consideración 1]
• [Consideración 2]

DEPENDENCIAS:
• [Dependencia 1]

IMPACTO: [Alto/Medio/Bajo]
COMPLEJIDAD: [Simple/Moderada/Compleja]

═══════════════════════════════════════════════════════════════
```
"""

    else:
        return create_prompt(document_text, role, 'HU Usuario Funcional')

    return prompt


def process_large_document(document_text, role, story_type, model, log_callback, progress_callback):
    """Procesa documentos grandes dividiéndolos en chunks."""
    try:
        log_callback("📄 Documento grande detectado. Iniciando análisis por fases...")

        # Fase 1: Análisis de funcionalidades
        log_callback("🔍 Fase 1: Identificando todas las funcionalidades...")
        analysis_prompt = create_analysis_prompt(document_text, role)

        progress_callback(20)
        analysis_response = model.generate_content(analysis_prompt, request_options={"timeout": 90})

        # Extraer lista de funcionalidades
        functionalities = []
        lines = analysis_response.text.split('\n')
        for line in lines:
            if re.match(r'^\d+\.', line.strip()):
                functionalities.append(line.strip())

        log_callback(f"✅ Identificadas {len(functionalities)} funcionalidades")

        # Fase 2: Generar historias por lotes
        all_stories = []
        batch_size = 5  # Procesar 5 funcionalidades a la vez
        total_batches = (len(functionalities) + batch_size - 1) // batch_size

        for batch_num in range(total_batches):
            start_idx = batch_num * batch_size
            log_callback(
                f"📝 Generando lote {batch_num + 1}/{total_batches} (funcionalidades {start_idx + 1}-{min(start_idx + batch_size, len(functionalities))})")

            story_prompt = create_story_generation_prompt(functionalities, document_text, role, start_idx, batch_size)

            progress = 30 + (batch_num / total_batches) * 60
            progress_callback(int(progress))

            try:
                story_response = model.generate_content(story_prompt, request_options={"timeout": 120})
                all_stories.append(story_response.text)
                log_callback(f"✅ Lote {batch_num + 1} completado")
            except Exception as e:
                log_callback(f"⚠️ Error en lote {batch_num + 1}: {e}")
                continue

        # Combinar todas las historias
        final_content = f"""
ANÁLISIS COMPLETO - {len(functionalities)} FUNCIONALIDADES IDENTIFICADAS
{"=" * 70}

FUNCIONALIDADES IDENTIFICADAS:
{chr(10).join(functionalities)}

{"=" * 70}
HISTORIAS DE USUARIO DETALLADAS
{"=" * 70}

{chr(10).join(all_stories)}

{"=" * 70}
RESUMEN FINAL
{"=" * 70}
✅ Total de funcionalidades procesadas: {len(functionalities)}
✅ Total de lotes generados: {total_batches}
✅ Análisis completado exitosamente
"""

        progress_callback(100)
        log_callback("🎉 Análisis completo finalizado exitosamente")

        return final_content

    except Exception as e:
        log_callback(f"❌ Error en procesamiento por chunks: {e}")
        return None


def generate_content_with_gemini(prompt, log_callback, progress_callback, model, document_text=None, role=None,
                                 story_type=None):
    """Envía el prompt a Gemini con manejo de documentos grandes."""
    try:
        # Detectar si necesita procesamiento por chunks
        if prompt == "CHUNK_PROCESSING_NEEDED" and document_text and role:
            return process_large_document(document_text, role, story_type, model, log_callback, progress_callback)

        log_callback("🤖 Generando contenido con Gemini...")
        progress_callback(30)

        # Para prompts normales
        response = model.generate_content(prompt, request_options={"timeout": 90})

        progress_callback(90)
        log_callback("✅ Generación completada")

        # Verificar si la respuesta se cortó
        response_text = response.text
        if "La generación completa" in response_text or "Este ejemplo ilustra" in response_text:
            log_callback("⚠️ Respuesta posiblemente incompleta detectada")

        return response_text

    except Exception as e:
        log_callback(f"❌ Error en la generación: {e}")
        return None


# --- FUNCIÓN PARA GUARDAR EN .DOCX ---
def save_as_docx(filepath, content):
    """Guarda el contenido en un archivo .docx."""
    try:
        doc = docx.Document()
        doc.add_paragraph(content)
        doc.save(filepath)
        return True
    except Exception as e:
        print(f"Error al guardar como DOCX: {e}")
        return False


# --- CLASE WORKER MEJORADA ---
class Worker(QObject):
    finished = Signal()
    progress_updated = Signal(int)
    log_message = Signal(str)

    def __init__(self, input_folder, output_folder, api_key, role, story_type):
        super().__init__()
        self.input_folder = input_folder
        self.output_folder = output_folder
        self.api_key = api_key
        self.role = role
        self.story_type = story_type

    def run(self):
        try:
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel("gemini-1.5-flash")

            files_to_process = [f for f in os.listdir(self.input_folder) if f.endswith(('.docx', '.pdf'))]
            total_files = len(files_to_process)

            if total_files == 0:
                self.log_message.emit("❌ No se encontraron archivos .docx o .pdf para procesar.")
                self.finished.emit()
                return

            for i, filename in enumerate(files_to_process):
                self.progress_updated.emit(int((i / total_files) * 100))
                file_path = os.path.join(self.input_folder, filename)
                self.log_message.emit(f"📂 Procesando: {filename}")

                text = extract_text_from_file(file_path)
                if not text.strip():
                    self.log_message.emit(f"⚠️ Archivo {filename} está vacío, saltando...")
                    continue

                self.log_message.emit(f"📄 Documento con {len(text)} caracteres")

                # Crear prompt y detectar si necesita procesamiento especial
                prompt = create_prompt(text, self.role, self.story_type)

                # Generar contenido (con manejo automático de chunks si es necesario)
                generated_text = generate_content_with_gemini(
                    prompt,
                    self.log_message.emit,
                    self.progress_updated.emit,
                    model,
                    document_text=text,
                    role=self.role,
                    story_type=self.story_type
                )

                if generated_text:
                    base_name = os.path.splitext(filename)[0]
                    output_file_path = os.path.join(self.output_folder, f"{base_name}_historias_completas.docx")

                    if save_as_docx(output_file_path, generated_text):
                        self.log_message.emit(f"✅ Archivo generado: {output_file_path}")
                    else:
                        self.log_message.emit(f"❌ Error al guardar: {output_file_path}")
                else:
                    self.log_message.emit(f"❌ No se pudo generar contenido para {filename}")

            self.progress_updated.emit(100)
            self.log_message.emit("🎉 ¡Proceso completado exitosamente!")

        except Exception as e:
            self.log_message.emit(f"💥 Error inesperado: {e}")
        finally:
            self.finished.emit()


# -----------------------------
# UI PySide6 (SIN CAMBIOS)
# -----------------------------
class GeneradorHistorias(QWidget):
    def __init__(self, stacked_widget):
        super().__init__()
        self.setWindowTitle("Generador de Historias de Usuario (IA)")
        self.resize(900, 750)
        self.stacked_widget = stacked_widget

        # --- Estilos Locales ---
        self.setStyleSheet("""
            QWidget {
                background-color: #1E1E2F;
                color: #F5F5F5;
                font-family: 'Segoe UI';
                font-size: 14px;
            }
            QFrame {
                background-color: #2A2A40;
                border-radius: 12px;
                border: 1px solid #3C3C5C;
            }
            QLabel {
                font-size: 16px;
                font-weight: bold;
                color: #B0C4DE;
            }
            QLineEdit, QTextEdit, QPlainTextEdit {
                background-color: #3A3A50;
                border: 1px solid #5C5C7A;
                color: #F5F5F5;
                border-radius: 5px;
                padding: 5px;
            }
            QComboBox {
                background-color: #3A3A50;
                border: 1px solid #5C5C7A;
                color: #F5F5F5;
                border-radius: 5px;
                padding: 5px;
            }
            QPushButton {
                background-color: #3F51B5;
                color: white;
                border-radius: 8px;
                padding: 10px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #5C6BC0;
            }
            QProgressBar {
                background-color: #2A2A40;
                border: 1px solid #3C3C5C;
                border-radius: 5px;
                text-align: center;
                color: #F5F5F5;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                border-radius: 5px;
            }
        """)

        layout = QVBoxLayout(self)

        # ---- Botón de Home ----
        btn_home = QPushButton("🏠 Menú Principal")
        btn_home.setFixedSize(150, 40)
        btn_home.setStyleSheet(
            "QPushButton { background-color: #3F51B5; color: white; border-radius: 8px; font-size: 14px; } QPushButton:hover { background-color: #5C6BC0; }")
        btn_home.clicked.connect(lambda: self.stacked_widget.setCurrentIndex(0))

        lbl_appname = QLabel("🚀 StoryCreator")
        lbl_appname.setStyleSheet("font-size: 18px; font-weight: bold; color: #FFD700;margin-right: 100px;")

        top_bar = QHBoxLayout()
        top_bar.addWidget(btn_home, alignment=Qt.AlignLeft)
        top_bar.addStretch()
        top_bar.addWidget(lbl_appname, alignment=Qt.AlignCenter)
        top_bar.addStretch()

        layout.addLayout(top_bar)

        # ---- Card configuración ----
        config_card = QFrame()
        config_card.setFrameShape(QFrame.StyledPanel)
        config_layout = QVBoxLayout(config_card)

        lbl_title = QLabel("⚙️ Configuración")
        lbl_title.setStyleSheet("font-size: 16px; font-weight: bold;")
        config_layout.addWidget(lbl_title)

        # Carpeta entrada
        row_in = QHBoxLayout()
        self.entry_input = QLineEdit()
        btn_in = QPushButton("Seleccionar")
        btn_in.clicked.connect(self.seleccionar_carpeta_in)
        row_in.addWidget(QLabel("Carpeta de Archivos:"))
        self.entry_input.setPlaceholderText(
            "Da click en el botón 'seleccionar' para colocar la ruta\n"
        )
        row_in.addWidget(self.entry_input)
        row_in.addWidget(btn_in)
        config_layout.addLayout(row_in)

        # Carpeta salida
        row_out = QHBoxLayout()
        self.entry_output = QLineEdit()
        btn_out = QPushButton("Seleccionar")
        btn_out.clicked.connect(self.seleccionar_carpeta_out)
        row_out.addWidget(QLabel("Carpeta de Salida:"))
        self.entry_output.setPlaceholderText(
            "Da click en el botón 'seleccionar' para colocar la ruta\n"
        )
        row_out.addWidget(self.entry_output)
        row_out.addWidget(btn_out)
        config_layout.addLayout(row_out)

        # API Key
        row_api = QHBoxLayout()
        self.entry_api = QLineEdit()
        self.entry_api.setEchoMode(QLineEdit.Password)
        row_api.addWidget(QLabel("API Key (Gemini):"))
        self.entry_api.setPlaceholderText(
            "Ejemplo: AIzaSyD3x-xxxx-xxxx-xxxx-xxxxx\n"
        )
        row_api.addWidget(self.entry_api)
        config_layout.addLayout(row_api)

        # Rol
        row_role = QHBoxLayout()
        self.cmb_role = QComboBox()
        self.cmb_role.addItems(["Usuario", "Administrador"])
        row_role.addWidget(QLabel("Rol:"))
        row_role.addWidget(self.cmb_role)
        config_layout.addLayout(row_role)

        # Tipo de historia
        row_story = QHBoxLayout()
        self.cmb_story = QComboBox()
        self.cmb_story.addItems(["HU Usuario Funcional", "HU no Funcionales", "Requisitos técnicos"])
        row_story.addWidget(QLabel("Tipo de Historia:"))
        row_story.addWidget(self.cmb_story)
        config_layout.addLayout(row_story)

        layout.addWidget(config_card)

        # ---- Acciones ----
        action_layout = QHBoxLayout()
        self.btn_run = QPushButton("🚀 Generar Historias")
        self.btn_run.clicked.connect(self.ejecutar)
        self.progress = QProgressBar()
        self.progress.setValue(0)

        action_layout.addWidget(self.btn_run)
        action_layout.addWidget(self.progress)
        layout.addLayout(action_layout)

        # ---- Log ----
        log_card = QFrame()
        log_card.setFrameShape(QFrame.StyledPanel)
        log_layout = QVBoxLayout(log_card)
        lbl_log = QLabel("📜 Resumen de ejecución")
        lbl_log.setStyleSheet("font-size: 14px; font-weight: bold;")
        log_layout.addWidget(lbl_log)

        self.txt_log = QPlainTextEdit()
        self.txt_log.setReadOnly(True)
        log_layout.addWidget(self.txt_log)

        layout.addWidget(log_card)

        # --- Atributos de QThread ---
        self.worker = None
        self.thread = None

    # ---- Funciones UI ----
    def seleccionar_carpeta_in(self):
        d = QFileDialog.getExistingDirectory(self, "Seleccionar carpeta de entrada")
        if d:
            self.entry_input.setText(d)

    def seleccionar_carpeta_out(self):
        d = QFileDialog.getExistingDirectory(self, "Seleccionar carpeta de salida")
        if d:
            self.entry_output.setText(d)

    def log_callback(self, msg):
        self.txt_log.appendPlainText(msg)

    def progress_callback(self, val):
        self.progress.setValue(val)

    def ejecutar(self):
        input_folder = self.entry_input.text().strip()
        output_folder = self.entry_output.text().strip()
        api_key = self.entry_api.text().strip()

        if not api_key:
            QMessageBox.critical(self, "Error", "Por favor ingresa tu API Key.")
            return
        if not input_folder or not output_folder:
            QMessageBox.critical(self, "Error", "Selecciona carpetas de entrada y salida.")
            return

        self.progress.setValue(0)
        self.log_callback("🚀 Iniciando proceso...")

        self.thread = QThread()
        self.worker = Worker(
            input_folder,
            output_folder,
            api_key,
            self.cmb_role.currentText(),
            self.cmb_story.currentText()
        )
        self.worker.moveToThread(self.thread)

        self.thread.started.connect(self.worker.run)
        self.worker.finished.connect(self.thread.quit)
        self.worker.finished.connect(self.worker.deleteLater)
        self.thread.finished.connect(self.thread.deleteLater)
        self.worker.log_message.connect(self.log_callback)
        self.worker.progress_updated.connect(self.progress_callback)

        self.thread.start()